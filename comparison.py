import streamlit as st
import pdfplumber
from docx import Document
from docx2pdf import convert
from difflib import Differ
import io
import platform  # Add this import



# Function to extract text from a PDF file
def extract_text_from_pdf(file):
    text = []
    with pdfplumber.open(file) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                paragraphs = page_text.split('\n')
                cleaned_paragraphs = []
                current_paragraph = ""
                for para in paragraphs:
                    cleaned_para = ' '.join(para.split())
                    if cleaned_para:
                        if current_paragraph:
                            if cleaned_para[0].islower() or para.startswith(' '):
                                current_paragraph += ' ' + cleaned_para
                            else:
                                cleaned_paragraphs.append((current_paragraph, "Normal"))
                                current_paragraph = cleaned_para
                        else:
                            current_paragraph = cleaned_para
                if current_paragraph:
                    cleaned_paragraphs.append((current_paragraph, "Normal"))
                text.extend(cleaned_paragraphs)
    return text

# Function to convert DOCX to PDF
def convert_docx_to_pdf(docx_path):
    pdf_path = docx_path.replace('.docx', '.pdf')
    convert(docx_path, pdf_path)
    return pdf_path

def convert_pdf_to_word(pdf_path, output_path):
    if platform.system() == "Windows":
        pythoncom.CoInitialize()
        word = win32com.client.Dispatch("Word.Application")
        doc = word.Documents.Open(pdf_path)
        doc.SaveAs(output_path, FileFormat=16)  # 16 represents the Word format
        doc.Close()
        word.Quit()
    else:
        raise NotImplementedError("PDF to Word conversion is only implemented for Windows.")

# Function to extract text from a DOCX file
def extract_text_from_docx(file):
    doc = Document(file)
    return [(para.text, "Normal") for para in doc.paragraphs if para.text]

# Function to load a document and extract text
def load_document(file):
    if file.name.endswith('.pdf'):
        return extract_text_from_pdf(file)
    elif file.name.endswith('.docx'):
        # Save the uploaded docx file to a temporary file
        with open(file.name, 'wb') as f:
            f.write(file.getvalue())
        pdf_path = convert_docx_to_pdf(file.name)
        return extract_text_from_pdf(pdf_path)
    return []

# Function to highlight differences between two texts
def highlight_differences(text1, text2):
    text1_str = text1.split()
    text2_str = text2.split()
    differ = Differ()
    diff = list(differ.compare(text1_str, text2_str))
    highlighted_text1 = []
    highlighted_text2 = []
    for word in diff:
        if word.startswith('? '):
            continue
        if word.startswith('- '):
            highlighted_text1.append(('red', word[2:]))
        elif word.startswith('+ '):
            highlighted_text2.append(('magenta', word[2:]))
        else:
            highlighted_text1.append(('black', word[2:] if word.startswith('  ') else word[2:]))
            highlighted_text2.append(('black', word[2:] if word.startswith('  ') else word[2:]))
    return highlighted_text1, highlighted_text2

# Function to render highlighted text as HTML with hover functionality
def render_highlighted_text(text, highlighted_text, editable=False):
    text_html = ""
    highlighted_words = iter(highlighted_text)
    continuous_text = ""
    current_tag = None
    for para_text in text.split('\n'):
        for word in para_text.split():
            try:
                tag, highlighted_word = next(highlighted_words)
                if tag != current_tag:
                    if continuous_text:
                        text_html += f'<span class="highlight {current_tag}" data-text="{continuous_text}">{continuous_text} </span>'
                    continuous_text = highlighted_word
                    current_tag = tag
                else:
                    continuous_text += ' ' + highlighted_word
            except StopIteration:
                if continuous_text:
                    text_html += f'<span class="highlight {current_tag}" data-text="{continuous_text}">{continuous_text} </span>'
                text_html += f'<span class="highlight black">{word} </span>'
                continuous_text = ""
                current_tag = None
        if continuous_text:
            text_html += f'<span class="highlight {current_tag}" data-text="{continuous_text}">{continuous_text} </span>'
            continuous_text = ""
            current_tag = None
        text_html += '<br>'
    
    # Add contenteditable attribute if required
    if editable:
        text_html = f'<div contenteditable="true" class="editable-container">{text_html}</div>'
    
    return text_html



# Initialize session state for document text storage
if 'doc1_text' not in st.session_state:
    st.session_state['doc1_text'] = None

if 'doc2_text' not in st.session_state:
    st.session_state['doc2_text'] = None

# Function to reset session state for new uploads
def reset_session_state():
    st.session_state['compare_clicked'] = False

# Streamlit application
st.set_page_config(page_title="COMPARE DOCUMENTS", page_icon="📄", layout="wide")

st.markdown(
    """
    <div class="header">
        <div class="logo-text">TEXT - GAZE</div>
    </div>
    """,
    unsafe_allow_html=True
)

st.markdown("""
    <script>
        document.addEventListener('DOMContentLoaded', function() {
            document.querySelectorAll('.editable-container').forEach(function(container) {
                container.addEventListener('input', function() {
                    let content = container.innerHTML;
                    container.setAttribute('data-content', content); // Store the content in a data attribute
                });
            });
        });

        document.addEventListener('DOMContentLoaded', function() {
            document.querySelectorAll('.highlight.red').forEach(function(elem) {
                elem.addEventListener('click', function(event) {
                    // Show options dialog on left-click
                    let options = confirm("Choose an option:\n\nOK - Delete\nCancel - Copy");
                    if (options) {
                        // Option selected - Delete text
                        elem.style.display = 'none';
                    } else {
                        // Option selected - Copy text
                        let range = document.createRange();
                        range.selectNode(elem);
                        window.getSelection().removeAllRanges();
                        window.getSelection().addRange(range);
                        document.execCommand('copy');
                        window.getSelection().removeAllRanges();
                        alert("Text copied to clipboard");
                    }
                });
            });

            document.querySelectorAll('.highlight.magenta').forEach(function(elem) {
                elem.addEventListener('click', function(event) {
                    // Handle moving magenta text to Document 1's container
                    let text = elem.getAttribute('data-text');
                    let doc1Container = document.querySelector('.editable-container');
                    let cursorPosition = window.getSelection().getRangeAt(0).startOffset;
                    let content = doc1Container.innerHTML;
                    let beforeCursor = content.slice(0, cursorPosition);
                    let afterCursor = content.slice(cursorPosition);
                    doc1Container.innerHTML = beforeCursor + '<span class="highlight magenta">' + text + '</span>' + afterCursor;
                });
            });
        });
    </script>

    """, unsafe_allow_html=True)




# Add custom CSS
st.markdown("""
    <style>
    body {
        background-color: #ffffff;
        margin: 0;
        padding: 0;
    }
    .editable-textarea {
        width: 100%;
        min-height: 500px; /* Set a minimum height */
        max-height: 90vh; /* Set a maximum height using viewport height */
        border: 1px solid #ccc;
        background-color: #f9f9f9;
        font-family: Arial, sans-serif;
        font-size: 14px;
        overflow: auto;
        padding: 10px;
        resize: vertical; /* Allow vertical resizing */
    }
    .editable-textarea::-webkit-scrollbar {
        width: 8px; /* Width of the scrollbar */
    }
    .editable-textarea::-webkit-scrollbar-track {
        background: #f1f1f1; /* Color of the track */
    }
    .editable-textarea::-webkit-scrollbar-thumb {
        background: #888; /* Color of the scroll thumb */
    }
    .editable-textarea::-webkit-scrollbar-thumb:hover {
        background: #555; /* Color of the scroll thumb on hover */
    }
    .editable-container {
    border: 1px solid #ccc;
    padding: 10px;
    min-height: 200px; /* Adjust as needed */
    }

    .highlight {
        position: relative;
    }
    .highlight::after {
        content: attr(data-text);
        position: absolute;
        background-color: #ffffe0; /* Light yellow background for better contrast */
        padding: 5px; /* Increase padding for readability */
        border: 1px solid #000; /* Solid black border */
        z-index: 1000; /* Increase z-index for better visibility */
        white-space: pre-wrap; /* Preserve whitespace and wrap text */
        display: none; /* Initially hidden */
        max-width: auto; /* Max-width set to auto */
        max-height: auto; /* Max-height set to auto */
        overflow: auto; /* Scrollbars for overflow content */
        font-size: 12px; /* Smaller font size for hover text */
        color: #000; /* Black text color */
        box-shadow: 0 2px 5px rgba(0, 0, 0, 0.3); /* Add shadow for better visibility */
        border-radius: 3px; /* Rounded corners for a softer look */
    }
    .highlight.red:hover::after,
    .highlight.magenta:hover::after,
    .highlight.blue:hover::after {
        display: block;
    }
    .highlight.red {color: red;}
    .highlight.magenta {color: magenta;}
    .highlight.blue {color: blue;}
    .highlight.black {color: black;}
    .button-center {
        height:50px
        position: absolute; /* Position the button absolutely within its container */
        bottom: 0; /* Stick it to the bottom of the column */
        left: 50%; /* Center horizontally */
        transform: translateX(-50%); /* Adjust for horizontal centering */
        margin-bottom: 0; /* Ensure no extra space at the top */
    }


    .center-text {
        display: flex;
        justify-content: center;
    }
    .header {
        width: 100%;
        height: 17%; /* Adjusted height for header */
        background-color: #dd6969; /* Reddish color for header */
        display: flex;
        align-items: center;
        justify-content: center;
        padding: 25px;
        box-shadow: 0 4px 8px rgba(0, 0, 0, 0.1);
        position: fixed; /* Fixed position */
        top: 0;
        left: 0;
        z-index: 1000;
    }
    .logo-text {
        font-size: 3.0rem; /* Increased font size */
        font-weight: bold; /* Bolder font */
        color: #fff;
        text-transform: uppercase; /* Text in capitals */
        text-align: center;
        width: 100%;
        position: absolute;
        height: 15%;
    }

    </style>
    <script>
    document.addEventListener('mousemove', function (e) {
        var highlights = document.querySelectorAll('.highlight:hover');
        highlights.forEach(function (highlight) {
            var tooltip = highlight.querySelector('::after');
            if (tooltip) {
                var rect = highlight.getBoundingClientRect();
                var tooltipHeight = tooltip.offsetHeight;
                var top = rect.top - tooltipHeight - 10;
                if (top < 0) {
                    top = rect.bottom + 10; // Move below the text if not enough space above
                }
                tooltip.style.top = top + 'px';
                tooltip.style.left = (rect.left + (rect.width / 2) - (tooltip.offsetWidth / 2)) + 'px';
            }
        });
    });
    </script>
    <script>
    document.addEventListener('DOMContentLoaded', function() {
        document.querySelectorAll('.editable-container').forEach(function(container) {
            container.addEventListener('input', function() {
             let content = container.innerHTML;
             container.setAttribute('data-content', content); // Store the content in a data attribute
            });
        });
    });
    </script>
    
        
            
    """, unsafe_allow_html=True)

# Create containers for document content
st.markdown('<h2 class="center-text">COMPARE DOCUMENTS</h2>', unsafe_allow_html=True)

# Create columns for document upload and comparison button
col1, col2, col3 = st.columns([1, 0.2, 1])

with col1:
    st.subheader("UPLOAD DOCUMENT 1")
    uploaded_file1 = st.file_uploader("CHOOSE A FILE (.docx), (.pdf)", type=["pdf", "docx"], key="doc1", on_change=reset_session_state)

    doc1_placeholder = st.empty()  # Placeholder for Document 1 content
    if uploaded_file1:
        if st.session_state['doc1_text'] is None:
            doc1_placeholder.text_area("DOCUMENT 1", "LOADING DOCUMENT 1...", height=300)
            doc1_placeholder.markdown("<h2 style='text-align: center;'>LOADING DOCUMENT 1</h2>", unsafe_allow_html=True)
            doc1_placeholder.markdown("""
                <style>
                .loader {
                    border: 16px solid #f3f3f3;
                    border-radius: 50%;
                    border-top: 16px solid #3498db;
                    width: 120px;
                    height: 120px;
                    animation: spin 2s linear infinite;
                    margin: auto;
                }

                @keyframes spin {
                    0% { transform: rotate(0deg); }
                    100% { transform: rotate(360deg); }
                }
                </style>
                <div class="loader"></div>
            """, unsafe_allow_html=True)
            st.session_state['doc1_text'] = load_document(uploaded_file1)
        doc1_text = "\n".join([para[0] for para in st.session_state['doc1_text']])
        edited_doc1_text = doc1_placeholder.text_area("Document 1", value=doc1_text, height=300)
        st.session_state['doc1_text'] = [(line, "Normal") for line in edited_doc1_text.split('\n')]


with col2:
    st.subheader(" ")
    if st.button("COMPARE", key="compare_button", help="Click to compare uploaded documents"):
        st.session_state['compare_clicked'] = True
    if st.button("SAVE(.docx)", key="docx_button", help="Save file as word format"):
        st.session_state['save_docx_button'] = True
    if st.button("SAVE(.pdf)", key="pdf_button", help="Save file as PDF format"):
        st.session_state['save_pdf_button'] = True

        


with col3:
    st.subheader("UPLOAD DOCUMENT 2")
    uploaded_file2 = st.file_uploader("CHOOSE A FILE (.docx), (.pdf)", type=["pdf", "docx"], key="doc2", on_change=reset_session_state)

    doc2_placeholder = st.empty()  # Placeholder for Document 2 content
    if uploaded_file2:
        if st.session_state['doc2_text'] is None:
            doc2_placeholder.text_area("DOCUMENT 2", "LOADING DOCUMENT 2...", height=300)
            doc2_placeholder.markdown("<h2 style='text-align: center;'>LOADING DOCUMENT 2</h2>", unsafe_allow_html=True)
            doc2_placeholder.markdown("""
                <style>
                .loader {
                    border: 16px solid #f3f3f3;
                    border-radius: 50%;
                    border-top: 16px solid #3498db;
                    width: 120px;
                    height: 120px;
                    animation: spin 2s linear infinite;
                    margin: auto;
                }

                @keyframes spin {
                    0% { transform: rotate(0deg); }
                    100% { transform: rotate(360deg); }
                }
                </style>
                <div class="loader"></div>
            """, unsafe_allow_html=True)
            st.session_state['doc2_text'] = load_document(uploaded_file2)
        doc2_text = "\n".join([para[0] for para in st.session_state['doc2_text']])
        doc2_placeholder.text_area("Document 2", value=doc2_text, height=300)

# Create column for comparison results
if st.session_state.get('compare_clicked', False):
    if st.session_state['doc1_text'] and st.session_state['doc2_text']:
        text1 = "\n".join([para[0] for para in st.session_state['doc1_text']])
        text2 = "\n".join([para[0] for para in st.session_state['doc2_text']])
        highlighted_text1, highlighted_text2 = highlight_differences(text1, text2)
        doc1_html = render_highlighted_text(text1, highlighted_text1, editable=True)
        doc2_html = render_highlighted_text(text2, highlighted_text2)

        # Replace the content of the placeholders with comparison results
        doc1_placeholder.empty()
        doc2_placeholder.empty()

        col1, col2 = st.columns([1, 1])
        with col1:
            st.subheader("DOCUMENT 1")
            st.markdown(f'<div class="editable-textarea">{doc1_html}</div>', unsafe_allow_html=True)
        with col2:
            st.subheader("DOCUMENT 2")
            st.markdown(f'<div class="editable-textarea">{doc2_html}</div>', unsafe_allow_html=True)
    else:
        st.error("!!!!PLEASE UPLOAD BOTH FILE BEFORE COMPARING!!!!")


# Function to get edited content from HTML (this requires JavaScript to be included in the HTML)
def get_edited_content():
    doc1_html = st.session_state.get('doc1_html', '')
    doc2_html = st.session_state.get('doc2_html', '')
    return doc1_html, doc2_html

if st.session_state.get('save_docx_button', False):
    # Retrieve the latest edited text from the session state
    edited_doc1_text = "\n".join([para[0] for para in st.session_state.get('doc1_text', [])])

    # Save the edited Document 1 as DOCX
    edited_doc1 = Document()
    edited_doc1.add_paragraph(edited_doc1_text)
    docx_buffer = io.BytesIO()
    edited_doc1.save(docx_buffer)
    docx_buffer.seek(0)
    st.download_button(label="DOWNLOAD DOCUMENT 1 (.docx)", data=docx_buffer, file_name="Document_1.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

if st.session_state.get('save_pdf_button', False):
    # Retrieve the latest edited text from the session state
    edited_doc1_text = "\n".join([para[0] for para in st.session_state.get('doc1_text', [])])

    # Save the edited Document 1 as PDF
    edited_doc1 = Document()
    edited_doc1.add_paragraph(edited_doc1_text)
    docx_path = "edited_document.docx"
    edited_doc1.save(docx_path)
    pdf_path = convert_docx_to_pdf(docx_path)
    with open(pdf_path, "rb") as pdf_file:
        st.download_button(label="DOWNLOAD DOCUMENT 1 (.pdf)", data=pdf_file, file_name="Document_1.pdf", mime="application/pdf")

